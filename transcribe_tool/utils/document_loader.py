"""
Document loading module for various text formats.

Supports loading text from multiple document formats:
- Plain text (.txt) - Best quality, direct text extraction
- PDF (.pdf) - Using pdfplumber for accurate extraction
- Word (.docx) - Using python-docx
- Markdown (.md) - Direct text with markup stripped
- RTF (.rtf) - Rich text format
- HTML (.html, .htm) - Web pages with tags stripped
- EPUB (.epub) - E-books
- ODT (.odt) - OpenDocument text
"""

import re
from pathlib import Path
from typing import List, Optional, Dict, Any, Union, Tuple
from dataclasses import dataclass, field
from enum import Enum
import logging


class DocumentFormat(Enum):
    """Supported document formats."""
    TXT = "txt"
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    MD = "md"
    RTF = "rtf"
    HTML = "html"
    HTM = "htm"
    EPUB = "epub"
    ODT = "odt"
    UNKNOWN = "unknown"


@dataclass
class Document:
    """Represents a loaded document with metadata."""
    content: str
    path: Optional[Path] = None
    format: DocumentFormat = DocumentFormat.UNKNOWN
    title: Optional[str] = None
    author: Optional[str] = None
    pages: int = 1
    word_count: int = 0
    char_count: int = 0
    language: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None
    success: bool = True

    def __post_init__(self):
        if self.content:
            self.word_count = len(self.content.split())
            self.char_count = len(self.content)


@dataclass
class LoaderResult:
    """Result of loading one or more documents."""
    documents: List[Document] = field(default_factory=list)
    total_files: int = 0
    successful: int = 0
    failed: int = 0
    errors: List[str] = field(default_factory=list)


class DocumentLoader:
    """
    Universal document loader supporting multiple text formats.

    Prioritizes high-quality text extraction with fallbacks
    for each format type.
    """

    # Supported extensions
    SUPPORTED_EXTENSIONS = {
        '.txt': DocumentFormat.TXT,
        '.pdf': DocumentFormat.PDF,
        '.docx': DocumentFormat.DOCX,
        '.doc': DocumentFormat.DOC,
        '.md': DocumentFormat.MD,
        '.markdown': DocumentFormat.MD,
        '.rtf': DocumentFormat.RTF,
        '.html': DocumentFormat.HTML,
        '.htm': DocumentFormat.HTM,
        '.epub': DocumentFormat.EPUB,
        '.odt': DocumentFormat.ODT,
    }

    def __init__(self, logger: Optional[logging.Logger] = None):
        self.logger = logger or logging.getLogger(__name__)
        self._check_dependencies()

    def _check_dependencies(self):
        """Check which optional dependencies are available."""
        self.has_pdfplumber = False
        self.has_pymupdf = False
        self.has_docx = False
        self.has_pypdf = False
        self.has_bs4 = False
        self.has_ebooklib = False
        self.has_odfpy = False
        self.has_striprtf = False

        try:
            import pdfplumber
            self.has_pdfplumber = True
        except ImportError:
            pass

        try:
            import fitz  # PyMuPDF
            self.has_pymupdf = True
        except ImportError:
            pass

        try:
            import pypdf
            self.has_pypdf = True
        except ImportError:
            try:
                import PyPDF2
                self.has_pypdf = True
            except ImportError:
                pass

        try:
            import docx
            self.has_docx = True
        except ImportError:
            pass

        try:
            from bs4 import BeautifulSoup
            self.has_bs4 = True
        except ImportError:
            pass

        try:
            import ebooklib
            self.has_ebooklib = True
        except ImportError:
            pass

        try:
            from odf import text as odf_text
            self.has_odfpy = True
        except ImportError:
            pass

        try:
            from striprtf.striprtf import rtf_to_text
            self.has_striprtf = True
        except ImportError:
            pass

    @classmethod
    def get_supported_formats(cls) -> List[str]:
        """Get list of supported file extensions."""
        return list(cls.SUPPORTED_EXTENSIONS.keys())

    def is_supported(self, path: Union[str, Path]) -> bool:
        """Check if a file format is supported."""
        path = Path(path)
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def get_format(self, path: Union[str, Path]) -> DocumentFormat:
        """Get the document format from file extension."""
        path = Path(path)
        return self.SUPPORTED_EXTENSIONS.get(path.suffix.lower(), DocumentFormat.UNKNOWN)

    def load(self, path: Union[str, Path]) -> Document:
        """
        Load a single document from file.

        Args:
            path: Path to the document file

        Returns:
            Document object with content and metadata
        """
        path = Path(path)

        if not path.exists():
            return Document(
                content="",
                path=path,
                success=False,
                error=f"File not found: {path}"
            )

        doc_format = self.get_format(path)

        try:
            if doc_format == DocumentFormat.TXT:
                return self._load_txt(path)
            elif doc_format == DocumentFormat.PDF:
                return self._load_pdf(path)
            elif doc_format in (DocumentFormat.DOCX, DocumentFormat.DOC):
                return self._load_docx(path)
            elif doc_format == DocumentFormat.MD:
                return self._load_markdown(path)
            elif doc_format == DocumentFormat.RTF:
                return self._load_rtf(path)
            elif doc_format in (DocumentFormat.HTML, DocumentFormat.HTM):
                return self._load_html(path)
            elif doc_format == DocumentFormat.EPUB:
                return self._load_epub(path)
            elif doc_format == DocumentFormat.ODT:
                return self._load_odt(path)
            else:
                # Try as plain text
                return self._load_txt(path)

        except Exception as e:
            self.logger.error(f"Error loading {path}: {e}")
            return Document(
                content="",
                path=path,
                format=doc_format,
                success=False,
                error=str(e)
            )

    def load_directory(
        self,
        directory: Union[str, Path],
        recursive: bool = True,
        extensions: Optional[List[str]] = None
    ) -> LoaderResult:
        """
        Load all supported documents from a directory.

        Args:
            directory: Path to directory
            recursive: Whether to search subdirectories
            extensions: Filter to specific extensions (e.g., ['.txt', '.pdf'])

        Returns:
            LoaderResult with list of Document objects
        """
        directory = Path(directory)
        result = LoaderResult()

        if not directory.exists():
            result.errors.append(f"Directory not found: {directory}")
            return result

        # Determine which extensions to look for
        if extensions:
            ext_set = set(ext.lower() if ext.startswith('.') else f'.{ext.lower()}'
                        for ext in extensions)
        else:
            ext_set = set(self.SUPPORTED_EXTENSIONS.keys())

        # Find files
        pattern = "**/*" if recursive else "*"
        files = [f for f in directory.glob(pattern)
                if f.is_file() and f.suffix.lower() in ext_set]

        result.total_files = len(files)

        for file_path in sorted(files):
            doc = self.load(file_path)
            result.documents.append(doc)

            if doc.success:
                result.successful += 1
            else:
                result.failed += 1
                if doc.error:
                    result.errors.append(f"{file_path}: {doc.error}")

        return result

    def load_text(self, text: str, source_name: str = "direct_input") -> Document:
        """
        Create a Document from raw text.

        Args:
            text: Raw text content
            source_name: Name to identify this text source

        Returns:
            Document object
        """
        return Document(
            content=self._clean_text(text),
            format=DocumentFormat.TXT,
            title=source_name,
            metadata={'source': 'direct_input'}
        )

    # =========================================================================
    # Format-specific loaders
    # =========================================================================

    def _load_txt(self, path: Path) -> Document:
        """Load plain text file."""
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']

        content = None
        used_encoding = None

        for encoding in encodings:
            try:
                content = path.read_text(encoding=encoding)
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            # Last resort: read as bytes and decode with errors ignored
            content = path.read_bytes().decode('utf-8', errors='ignore')
            used_encoding = 'utf-8 (with errors ignored)'

        return Document(
            content=self._clean_text(content),
            path=path,
            format=DocumentFormat.TXT,
            title=path.stem,
            metadata={'encoding': used_encoding}
        )

    def _load_pdf(self, path: Path) -> Document:
        """Load PDF document with best available library."""
        content = ""
        pages = 0
        metadata = {}

        # Try pdfplumber first (best quality)
        if self.has_pdfplumber:
            try:
                import pdfplumber
                text_parts = []

                with pdfplumber.open(path) as pdf:
                    pages = len(pdf.pages)
                    metadata = pdf.metadata or {}

                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)

                content = "\n\n".join(text_parts)

                if content.strip():
                    return Document(
                        content=self._clean_text(content),
                        path=path,
                        format=DocumentFormat.PDF,
                        title=metadata.get('Title') or path.stem,
                        author=metadata.get('Author'),
                        pages=pages,
                        metadata={'extractor': 'pdfplumber', **metadata}
                    )
            except Exception as e:
                self.logger.debug(f"pdfplumber failed: {e}")

        # Try PyMuPDF (fast and accurate)
        if self.has_pymupdf:
            try:
                import fitz
                text_parts = []

                doc = fitz.open(path)
                pages = len(doc)
                metadata = doc.metadata or {}

                for page in doc:
                    text_parts.append(page.get_text())

                doc.close()
                content = "\n\n".join(text_parts)

                if content.strip():
                    return Document(
                        content=self._clean_text(content),
                        path=path,
                        format=DocumentFormat.PDF,
                        title=metadata.get('title') or path.stem,
                        author=metadata.get('author'),
                        pages=pages,
                        metadata={'extractor': 'pymupdf', **metadata}
                    )
            except Exception as e:
                self.logger.debug(f"PyMuPDF failed: {e}")

        # Try pypdf/PyPDF2
        if self.has_pypdf:
            try:
                try:
                    from pypdf import PdfReader
                except ImportError:
                    from PyPDF2 import PdfReader

                reader = PdfReader(path)
                pages = len(reader.pages)
                text_parts = []

                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                content = "\n\n".join(text_parts)

                if content.strip():
                    return Document(
                        content=self._clean_text(content),
                        path=path,
                        format=DocumentFormat.PDF,
                        title=path.stem,
                        pages=pages,
                        metadata={'extractor': 'pypdf'}
                    )
            except Exception as e:
                self.logger.debug(f"pypdf failed: {e}")

        # No PDF library available or all failed
        if not (self.has_pdfplumber or self.has_pymupdf or self.has_pypdf):
            return Document(
                content="",
                path=path,
                format=DocumentFormat.PDF,
                success=False,
                error="No PDF library available. Install: pip install pdfplumber"
            )

        return Document(
            content=self._clean_text(content) if content else "",
            path=path,
            format=DocumentFormat.PDF,
            title=path.stem,
            pages=pages,
            success=bool(content.strip()),
            error="Could not extract text from PDF" if not content.strip() else None
        )

    def _load_docx(self, path: Path) -> Document:
        """Load Word document (.docx)."""
        if not self.has_docx:
            return Document(
                content="",
                path=path,
                format=DocumentFormat.DOCX,
                success=False,
                error="python-docx not installed. Install: pip install python-docx"
            )

        try:
            import docx

            doc = docx.Document(path)
            text_parts = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            content = "\n\n".join(text_parts)

            # Try to get metadata
            core_props = doc.core_properties

            return Document(
                content=self._clean_text(content),
                path=path,
                format=DocumentFormat.DOCX,
                title=core_props.title or path.stem,
                author=core_props.author,
                metadata={
                    'subject': core_props.subject,
                    'keywords': core_props.keywords,
                    'created': str(core_props.created) if core_props.created else None,
                    'modified': str(core_props.modified) if core_props.modified else None,
                }
            )
        except Exception as e:
            return Document(
                content="",
                path=path,
                format=DocumentFormat.DOCX,
                success=False,
                error=str(e)
            )

    def _load_markdown(self, path: Path) -> Document:
        """Load Markdown file and optionally strip markup."""
        content = path.read_text(encoding='utf-8')

        # Strip markdown formatting for cleaner text
        clean_content = self._strip_markdown(content)

        return Document(
            content=self._clean_text(clean_content),
            path=path,
            format=DocumentFormat.MD,
            title=path.stem,
            metadata={'raw_markdown': content}
        )

    def _load_rtf(self, path: Path) -> Document:
        """Load RTF file."""
        if not self.has_striprtf:
            return Document(
                content="",
                path=path,
                format=DocumentFormat.RTF,
                success=False,
                error="striprtf not installed. Install: pip install striprtf"
            )

        try:
            from striprtf.striprtf import rtf_to_text

            rtf_content = path.read_text(encoding='utf-8', errors='ignore')
            content = rtf_to_text(rtf_content)

            return Document(
                content=self._clean_text(content),
                path=path,
                format=DocumentFormat.RTF,
                title=path.stem
            )
        except Exception as e:
            return Document(
                content="",
                path=path,
                format=DocumentFormat.RTF,
                success=False,
                error=str(e)
            )

    def _load_html(self, path: Path) -> Document:
        """Load HTML file and extract text."""
        content = path.read_text(encoding='utf-8', errors='ignore')

        if self.has_bs4:
            try:
                from bs4 import BeautifulSoup

                soup = BeautifulSoup(content, 'html.parser')

                # Remove script and style elements
                for script in soup(['script', 'style', 'head', 'meta', 'link']):
                    script.decompose()

                # Get title
                title = soup.title.string if soup.title else path.stem

                # Get text
                text = soup.get_text(separator='\n')

                return Document(
                    content=self._clean_text(text),
                    path=path,
                    format=DocumentFormat.HTML,
                    title=title
                )
            except Exception as e:
                self.logger.debug(f"BeautifulSoup failed: {e}")

        # Fallback: simple regex-based tag stripping
        text = self._strip_html_tags(content)

        return Document(
            content=self._clean_text(text),
            path=path,
            format=DocumentFormat.HTML,
            title=path.stem
        )

    def _load_epub(self, path: Path) -> Document:
        """Load EPUB e-book."""
        if not self.has_ebooklib:
            return Document(
                content="",
                path=path,
                format=DocumentFormat.EPUB,
                success=False,
                error="ebooklib not installed. Install: pip install ebooklib"
            )

        try:
            import ebooklib
            from ebooklib import epub

            book = epub.read_epub(path)
            text_parts = []

            # Get metadata
            title = book.get_metadata('DC', 'title')
            title = title[0][0] if title else path.stem

            author = book.get_metadata('DC', 'creator')
            author = author[0][0] if author else None

            # Extract text from each item
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    content = item.get_content().decode('utf-8', errors='ignore')

                    if self.has_bs4:
                        from bs4 import BeautifulSoup
                        soup = BeautifulSoup(content, 'html.parser')
                        text_parts.append(soup.get_text(separator='\n'))
                    else:
                        text_parts.append(self._strip_html_tags(content))

            content = "\n\n".join(text_parts)

            return Document(
                content=self._clean_text(content),
                path=path,
                format=DocumentFormat.EPUB,
                title=title,
                author=author
            )
        except Exception as e:
            return Document(
                content="",
                path=path,
                format=DocumentFormat.EPUB,
                success=False,
                error=str(e)
            )

    def _load_odt(self, path: Path) -> Document:
        """Load OpenDocument text file."""
        if not self.has_odfpy:
            return Document(
                content="",
                path=path,
                format=DocumentFormat.ODT,
                success=False,
                error="odfpy not installed. Install: pip install odfpy"
            )

        try:
            from odf import text as odf_text
            from odf.opendocument import load

            doc = load(path)
            text_parts = []

            for paragraph in doc.getElementsByType(odf_text.P):
                text = ""
                for node in paragraph.childNodes:
                    if hasattr(node, 'data'):
                        text += node.data
                if text.strip():
                    text_parts.append(text)

            content = "\n\n".join(text_parts)

            return Document(
                content=self._clean_text(content),
                path=path,
                format=DocumentFormat.ODT,
                title=path.stem
            )
        except Exception as e:
            return Document(
                content="",
                path=path,
                format=DocumentFormat.ODT,
                success=False,
                error=str(e)
            )

    # =========================================================================
    # Text cleaning utilities
    # =========================================================================

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""

        # Normalize unicode
        import unicodedata
        text = unicodedata.normalize('NFKC', text)

        # Replace various dash types with standard hyphen
        text = re.sub(r'[\u2010-\u2015\u2212]', '-', text)

        # Replace various quote types
        text = re.sub(r'[\u2018\u2019\u201A\u201B]', "'", text)
        text = re.sub(r'[\u201C\u201D\u201E\u201F]', '"', text)

        # Normalize whitespace
        text = re.sub(r'\t', ' ', text)
        text = re.sub(r' +', ' ', text)

        # Normalize line endings
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)

        # Remove excessive blank lines (keep max 2)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Strip leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)

        return text.strip()

    def _strip_html_tags(self, html: str) -> str:
        """Simple HTML tag removal without BeautifulSoup."""
        # Remove script and style content
        html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL | re.IGNORECASE)

        # Replace block elements with newlines
        html = re.sub(r'<(p|div|br|h[1-6]|li|tr)[^>]*>', '\n', html, flags=re.IGNORECASE)

        # Remove all remaining tags
        html = re.sub(r'<[^>]+>', '', html)

        # Decode HTML entities
        html = re.sub(r'&nbsp;', ' ', html)
        html = re.sub(r'&lt;', '<', html)
        html = re.sub(r'&gt;', '>', html)
        html = re.sub(r'&amp;', '&', html)
        html = re.sub(r'&quot;', '"', html)
        html = re.sub(r'&#(\d+);', lambda m: chr(int(m.group(1))), html)

        return html

    def _strip_markdown(self, text: str) -> str:
        """Remove Markdown formatting from text."""
        # Remove code blocks
        text = re.sub(r'```[\s\S]*?```', '', text)
        text = re.sub(r'`[^`]+`', '', text)

        # Remove headers
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)

        # Remove emphasis
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'__([^_]+)__', r'\1', text)
        text = re.sub(r'_([^_]+)_', r'\1', text)

        # Remove links but keep text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

        # Remove images
        text = re.sub(r'!\[[^\]]*\]\([^)]+\)', '', text)

        # Remove horizontal rules
        text = re.sub(r'^[-*_]{3,}\s*$', '', text, flags=re.MULTILINE)

        # Remove list markers
        text = re.sub(r'^[\s]*[-*+]\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'^[\s]*\d+\.\s+', '', text, flags=re.MULTILINE)

        # Remove blockquotes
        text = re.sub(r'^>\s*', '', text, flags=re.MULTILINE)

        return text


def load_document(path: Union[str, Path]) -> Document:
    """
    Convenience function to load a single document.

    Args:
        path: Path to document file

    Returns:
        Document object
    """
    loader = DocumentLoader()
    return loader.load(path)


def load_documents(
    directory: Union[str, Path],
    recursive: bool = True,
    extensions: Optional[List[str]] = None
) -> List[Document]:
    """
    Convenience function to load all documents from a directory.

    Args:
        directory: Path to directory
        recursive: Search subdirectories
        extensions: Filter by extensions

    Returns:
        List of Document objects
    """
    loader = DocumentLoader()
    result = loader.load_directory(directory, recursive, extensions)
    return result.documents
